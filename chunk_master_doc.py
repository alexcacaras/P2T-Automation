"""
Chunk Master Document into Per-Task Knowledge Files
====================================================
Reads the P2T Master Document (.docx) and splits it into JSON chunks
by detecting Word heading styles. Each chunk maps to an automation task.

USAGE:
    python chunk_master_doc.py                           # uses default path
    python chunk_master_doc.py path/to/master_doc.docx   # custom path

OUTPUT:
    knowledge/task_0_pretask_setup.json
    knowledge/task_1_disable_notifications.json
    knowledge/task_3_disable_adp_deliveries.json
    ... etc
"""

import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# Heading text → task mapping (handles both old and new heading formats)
HEADING_MAP = [
    
    ("Disable or Update ADP", 3, "Disable ADP Deliveries", "task3_disable_adp_deliveries"),
    ("Remove Receivables System Email", 12, "Remove Receivables Emails", "task12_remove_receivables_emails"),
    ("Required Roles and Access", 0, "Pretask - Procurement Access Setup", "setup_procurement_access_for_user"),
    ("Check if Notifications are set to", 1, "Disable Notifications", "task1_disable_notifications"),
    ("Change Color Scheme and Banner Message", 2, "Banner Message and Color Scheme", "task2_banner_message"),
    ("Banner Message", 2, "Banner Message", "task2_banner_message"),
    ("Disable/Update ADP Integrations", 3, "Disable ADP Deliveries", "task3_disable_adp_deliveries"),
    ("Add IP addresses", 4, "Add IP Addresses", "task4_add_ip"),
    ("Turn off PO Communication", 7, "Turn Off PO Communication", "task7_turn_off_po_communication"),
    ("Disabling AP Payment Transmission", 9, "Disable AP Payment Transmission", "task9_disable_ap_payment"),
    ("Update Corp Card Program", 10, "Update Corp Card Program", "task10_update_corp_card"),
    ("Disable GetThere Configuration", 11, "Disable GetThere Configuration", "task11_disable_getthere"),
    ("Remove the (email from) receivable", 12, "Remove Receivables Emails", "task12_remove_receivables_emails"),
    ("Update or Remove HireRight", 15, "Update HireRight Configuration", "task15_update_hireright"),
    ("Pre-Note", 16, "PreNote Update SFTP / Disable Delivery", "task16_prenote"),
    ("ADMIN User Accounts creation", 17, "Admin User Accounts Creation", "task17_admin_users"),
    ("User Accounts for API", 18, "User Accounts for API (Integration)", "task18_tech_user"),
    ("Disable Separate Remittance Advice", 21, "Disable Separate Remittance Emails", "task21_disable_remittance"),
    ("Journeys tasks with external URL", 22, "Journey Tasks with External URL", "task22_update_checklist_urls"),
    ("Search for Leave of Absence", 22, "Journey - Leave of Absence", "task22_update_checklist_urls"),
    ("Medical Leave journey", 22, "Journey - Medical Leave", "task22_update_checklist_urls"),
    ("Leave Extension", 22, "Journey - Leave Extension", "task22_update_checklist_urls"),
    ("Work Force Structure: Positions", 23, "Workforce Structure Positions", "task23_workforce_structure"),
]


def match_heading(text):
    """Match a paragraph's text to a heading in our map."""
    for heading, num, name, code_fn in HEADING_MAP:
        if heading.lower() in text.lower():
            return {"num": num, "name": name, "code_fn": code_fn, "heading": heading}
    return None


def extract_expected_elements(text):
    """Extract UI element names mentioned in the steps."""
    elements = set()
    for match in re.finditer(r'["\u201c]([^"\u201d]{3,60})["\u201d]', text):
        elements.add(match.group(1).strip())
    for kw in ["Navigator", "Setup and Maintenance", "Security Console",
               "Schedule New Process", "Save and Close", "Save", "Done",
               "Submit", "Search", "Show Filters", "Validate",
               "Data Exchange", "Extract Definitions", "Deliver",
               "Additional Details", "Report Categories", "Worklist",
               "Administration", "My Client Groups", "Manage Administrator Profile Values",
               "hamburger", "three lines", "navigation"]:
        if kw.lower() in text.lower():
            elements.add(kw)
    return sorted(elements)


def chunk_document(docx_path, output_dir):
    print(f"Reading: {docx_path}")
    doc = Document(str(docx_path))
    
    output_dir = Path(output_dir)
    output_dir.mkdir(exist_ok=True, parents=True)
    
    # ── Step 0: Extract pretask section (everything before first Head 1) ──
    pretask_lines = []
    first_heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if para.style.name in ('Heading 2', 'Head 1 Novamodus', 'Head 2 Novamodus'):
            first_heading_idx = i
            break
        if 'TOC Heading' in para.style.name or i >= 6:  # Skip cover page
            pretask_lines.append(text)
    
    if pretask_lines:
        pretask_text = "\n".join(pretask_lines)
        pretask_chunk = {
            "task_num": 0,
            "task_name": "Pretask - Procurement Access Setup",
            "code_function": "setup_procurement_access_for_user",
            "doc_section": "Required roles and data access setup",
            "steps": pretask_text,
            "expected_elements": extract_expected_elements(pretask_text),
        }
        fp = output_dir / "task_0_pretask_setup.json"
        fp.write_text(json.dumps(pretask_chunk, indent=2, ensure_ascii=False), encoding="utf-8")
        print(f"  ✓ task_0_pretask_setup.json (Task 0: Pretask Setup)")
    
    # ── Step 1: Find all heading sections ──
    sections = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        if para.style.name in ('Heading 2', 'Head 1 Novamodus', 'Head 2 Novamodus'):
            match = match_heading(text)
            if match:
                sections.append((i, match, [text]))
    
    print(f"Found {len(sections)} heading sections + pretask")
    
    # ── Step 2: Assign content paragraphs to each section ──
    for idx, (para_idx, match, content) in enumerate(sections):
        next_para_idx = sections[idx + 1][0] if idx + 1 < len(sections) else len(doc.paragraphs)
        for j in range(para_idx + 1, next_para_idx):
            text = doc.paragraphs[j].text.strip()
            if text:
                content.append(text)
    
    # ── Step 3: Build and save JSON chunks ──
    created_files = {"task_0_pretask_setup.json": True}
    for para_idx, match, content in sections:
        section_text = "\n".join(content)
        
        chunk = {
            "task_num": match["num"],
            "task_name": match["name"],
            "code_function": match["code_fn"],
            "doc_section": match["heading"],
            "steps": section_text,
            "expected_elements": extract_expected_elements(section_text),
        }
        
        safe_name = re.sub(r'[^a-z0-9]+', '_', match["name"].lower()).strip('_')
        filename = f"task_{match['num']}_{safe_name}.json"
        filepath = output_dir / filename
        
        # Merge if same file already exists (e.g. Task 22 has multiple journey sections)
        if filename in created_files:
            existing = json.loads(filepath.read_text(encoding="utf-8"))
            existing["steps"] += "\n\n--- Additional Section ---\n\n" + section_text
            existing["expected_elements"] = sorted(
                set(existing["expected_elements"]) | set(chunk["expected_elements"])
            )
            chunk = existing
        
        filepath.write_text(json.dumps(chunk, indent=2, ensure_ascii=False), encoding="utf-8")
        created_files[filename] = True
        print(f"  ✓ {filename} (Task {match['num']}: {match['name']})")
    
    # ── Step 4: Combined reference ──
    all_chunks = []
    for f in sorted(output_dir.glob("task_*.json")):
        all_chunks.append(json.loads(f.read_text(encoding="utf-8")))
    (output_dir / "_all_tasks.json").write_text(
        json.dumps(all_chunks, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    
    print(f"\nDone! {len(created_files)} chunk files + _all_tasks.json in {output_dir}/")


if __name__ == "__main__":
    doc_path = Path(sys.argv[1]) if len(sys.argv) > 1 else None
    if not doc_path:
        for c in [Path("knowledge/P2T_Mater_Document-_(client).docx"), Path("P2T_Mater_Document-_(client).docx")]:
            if c.exists():
                doc_path = c
                break
    if not doc_path:
        print("Usage: python chunk_master_doc.py path/to/master_doc.docx")
        sys.exit(1)
    
    chunk_document(doc_path, Path(__file__).parent / "knowledge")